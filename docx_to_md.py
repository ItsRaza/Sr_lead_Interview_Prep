"""Convert Word .docx files to Markdown by extracting paragraphs and headings."""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("python-docx not installed. Run: python -m pip install python-docx")
    exit(1)

def get_style_name(paragraph):
    """Get style name if it's a heading or normal."""
    if paragraph.style is None:
        return None
    name = (paragraph.style.name or "").lower()
    if "heading" in name or name.startswith("title"):
        return name
    return "normal"

def heading_level(style_name):
    """Map Word heading style to markdown level (1-6)."""
    if not style_name:
        return 0
    name = style_name.lower()
    if "title" in name:
        return 1
    for i in range(1, 10):
        if f"heading {i}" in name or f"heading{i}" in name:
            return min(i, 6)
    if "heading" in name:
        return 2
    return 0

def paragraph_to_md(paragraph):
    """Convert a single paragraph to markdown line(s)."""
    text = paragraph.text.strip()
    if not text:
        return ""
    style_name = get_style_name(paragraph)
    level = heading_level(style_name)
    if level >= 1:
        return "#" * level + " " + text
    # Bold at start often indicates a label
    return text

def docx_to_md(docx_path):
    """Read docx and return markdown string."""
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        line = paragraph_to_md(para)
        if line:
            lines.append(line)
    # Handle tables: simple markdown tables
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            lines.append("")
            lines.extend(rows)
            if len(rows) >= 2:
                # Add separator after header
                num_cols = len(table.rows[0].cells)
                lines.insert(-len(rows) + 1, "|" + " --- |" * num_cols)
            lines.append("")
    return "\n\n".join(lines)

def main():
    base = Path(__file__).parent
    for name in [
        "Senior_Lead_Architecture_Interview_Prep_Part1.docx",
        "Senior_Lead_Architecture_Interview_Prep_Part2.docx",
        "Senior_Lead_Architecture_Interview_Prep_Part3.docx",
    ]:
        docx_path = base / name
        if not docx_path.exists():
            print(f"Skip (not found): {name}")
            continue
        md_name = name.replace(".docx", ".md")
        md_path = base / md_name
        md_content = docx_to_md(docx_path)
        md_path.write_text(md_content, encoding="utf-8")
        print(f"Created: {md_name}")

if __name__ == "__main__":
    main()
